#########################
#Russell Arlt           #
#CIT 112 David Hosler   #
#Multiplication Quiz    #
#########################

import sys
import pyinputplus as pyip
import random as rand
import time
import datetime
import docx # importing all these

def main(human_name): #starting a main function
    max_multi = pyip.inputInt(f"{human_name} what is the max multiplier you want to see?: ", min=1, max=999) #making variable max_multi take only an integer input, at a min of 1 and max of 9
    quest_num = pyip.inputInt(f"{human_name} how many questions would you like to answer?", min=1, max=999)  # setting variable equal to the input value of another integer. If not repeats.
    input("Press enter to start the quiz!")  # lets start this quiz.
    save_time = datetime.datetime.now() # get the time right now and save it to a variable
    fun_time = save_time.strftime("%Y%m%d%H%M") # convert the variable to Year, numerical month, day, hour and minute.
    save_it_up_bb = (f'{human_name}' + '_' + f'{fun_time}') # Saving the name and date time format as required.
    doc = docx.Document() # Create a document
    doc.add_heading(f'Header {human_name}\'s Multiplication Quiz', 0) # Create a heading
    correct, incorrect, question_num = questions(quest_num, max_multi, doc, save_it_up_bb) #The variables correct and incorrect will be equal to the questions where we are taking the quest num and max multi
    stats(correct, incorrect, doc, save_it_up_bb, question_num) #passing correct and incorrect into the stats function
    play_again(human_name) #calling the play again function

def questions(quest_num, max_multi, doc, save_it_up_bb): #creation of the function questions, taking in quest num and max multi
    correct = 0 #starting variable correct at 0
    incorrect = 0 #starting variable incorrect at 0
    question_num = 1 # a numerical variable set at one that increments based on how many times the loop loops.
    while True: # loop
        for num in range(0, quest_num): #for variable in range 0 to the number of questions to answer.
            num_random_one = rand.randint(0, max_multi) #creating a random number between 0 and the max multi number declared
            num_random_two = rand.randint(0, max_multi) #creating a second random number between 0 and the max multi number declared
            var = (f"what is {num_random_one} x {num_random_two}?") #asking the user to do math!
            print(var) # print that question
            doc.add_paragraph(f'Question #{question_num}') # add a paragraph to the doc showing the question number
            doc.add_paragraph(f"{var}") # add another paragraph  showing the question
            the_number = num_random_one * num_random_two  # creating the answer from the random numbers
            tries = 0 #variable for the number of tries they will have.
            while True: #new loop
                try: # lets try this
                    answer = pyip.inputInt("answer: ", timeout=10, limit=3) #the answer they supply must be an int and we set a timeout of 10 and limit 3
                    doc.add_paragraph(f'{human_name}\'s answer is: {answer}\n') #putting the users answer in the soc.
                except pyip.RetryLimitException: #if this exception is thrown
                    print("limit of 3 reached, question incorrect.") #What is said if thats thrown
                    print(f"correct answer is: {the_number}") #shows the correct number
                    incorrect += 1 #increases the incorrect variable
                    question_num += 1 # increment
                    break #breaks the loop
                except pyip.TimeoutException: #if this exception is thrown
                    print("timeout, question incorrect.") #print this
                    print(f"correct answer is: {the_number}")  # shows the correct number
                    incorrect += 1 #increment incorrect variable by 1
                    question_num += 1 # increment
                    break #break
                if answer == the_number: # if the answer the user gives is equal to the number
                    print("Correct! You is smart and good.") #print this
                    correct += 1 #increment correct variable by 1
                    question_num += 1 # increment
                    break #break
                elif answer is not the_number:# if they don't input the correct answer
                    incorrect += 1 #increment incorrect 1
                    tries += 1 #increment tries 1
                    if tries < 3: #count the number of tries, if its less than 3
                        print("wrong, try again") #let them try again!
                    elif tries >= 3: # count the number of tries if it is more than or equal to 3
                        print("correct attempt limit reached.") # they didn't get it in time!
                        print(f"correct answer is: {the_number}") #show them the folly of their ways
                        question_num += 1 # increment
                        break #break it down!
                    continue #continue back up!
                else: #if somehow you get here, not sure how you would.
                    print("secret zone!") #easter egg
                    continue #keep going
        return correct, incorrect, question_num #return both correct and incorrect variables and the number of questions

def stats(correct, incorrect, doc, save_it_up_bb, question_num): #defining the stats function that is taking correct and incorrect
    toc = time.time() #takes the time at the moment that the user is in this spot.
    newtime = toc - tic # subtracts the original time from toc
    print("lets look at the stats!") #lets see these stats!
    var1 = f"You played {question_num} games!" # the number of games set to a variable
    print(var1) # printed variable
    print(f"You played for {newtime} seconds!") #shows the user how many seconds they played for.
    var2 = f"you got the answer correct, {correct} times! " # variable set to the number of times they were correct
    print(var2) # shows the user how many you got correct
    print(f"you got the answer wrong, {incorrect} times.") #shows the user how many you got incorrect
    new_value = correct+incorrect # creates a new value from the correct and incorrect variables
    try: #try this!
        calc_percentage = (correct/new_value) * 100 #calculating the percentage by dividing correct by the new value variable and multiplying by 100
    except ZeroDivisionError: # If we get a zero division error
        if correct == 0: # and correct is 0
            calc_percentage = 0 #we will make the percentage at 0
        elif incorrect == 0: # or incorrect is 0
            calc_percentage = 100 #we will set percentage at 100
    var3 = f"your percentage right is {calc_percentage}%" # variable set to the percentage correct
    section = doc.sections[0] # setting a variable named sections to doc.sections
    footer = section.footer # varaible set to the footer section
    p = footer.paragraphs[0] # variable set to paragraphs of the footer section
    p.text = (f'{var1}\n{var2}\n{var3}\n') # putting text into the paragraph equal to the string
    doc.save(f'{save_it_up_bb}.docx') # save the doc.
    print(var3) # print the percentage

def play_again(human_name): #define the function play again
    print("play again?") # ask them if they want to play again
    play_again = pyip.inputYesNo() #using the pyinputplus yes or no function to get a yes or no from the user
    if play_again == "yes": # if they do get a yes back from input yes no
        main(human_name) # calling back to the main function
    else: #else
        print("ok bye bye") #a sweet goodbye for you
        sys.exit() #leave the program

tic = time.time() #first part of the program, starting a timer
human_name = input("Hello human, what do other humans call you?: ")
if human_name == '': # if they put nothing
    print('We will call you Hugh Man then.') # they become hugh man
    human_name = 'Hugh Man' # thats right
main(human_name) #calling us to the main function right away.